"""SPK document generation. Builds the .docx in memory and hands back the
stream — nothing here touches the filesystem."""

import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import renderer
from terbilang import terbilang

# Letterhead. The one place these appear — change them here and every SPK
# issued afterwards carries the new details.
KOP_NAMA = "PT Nusantara Kreasi Acara"
KOP_ALAMAT = "Jl. Gatot Subroto No. 42, Jakarta Selatan 12190"
KOP_TELEPON = "Telp. (021) 555-0142"
KOP_EMAIL = "procurement@nusantarakreasi.co.id"
# Sits before the date on the signature block: "Jakarta, 10 Agustus 2026".
KOP_KOTA = "Jakarta"

FONT_NAMA = "Times New Roman"
FONT_UKURAN = Pt(11)


def format_rupiah(harga: int, prefix: bool = True) -> str:
    """15000000 -> "Rp 15.000.000", or "15.000.000" without the prefix, which
    is the shape the SPK form's input reads back. Thousands separated with
    dots, the Indonesian convention, and no decimals — harga is whole rupiah."""
    angka = f"{harga:,}".replace(",", ".")
    return f"Rp {angka}" if prefix else angka


def slug(teks: str) -> str:
    """Vendor name as a filename part."""
    bersih = re.sub(r"[^a-z0-9]+", "-", str(teks).lower()).strip("-")
    return bersih or "vendor"


def nama_berkas_spk(nomor: str, nama_pt: str) -> str:
    """The document's filename. It lives here because it is a property of the
    document, not of the route that happens to serve it — the slashes in a
    nomor are path separators, so they become dashes."""
    return f"SPK-{str(nomor).replace('/', '-')}-{slug(nama_pt)}.docx"


# Everything buat_spk_docx needs that the SPK form does not itself collect.
# The form checks this list before offering to issue, so a raise from here
# never surfaces as a 500 — keep it in step with the _teks calls below.
KONTEKS_WAJIB = (
    ("vendor", "nama_pt", "vendor company name"),
    ("vendor", "pic_nama", "vendor contact name (PIC)"),
    ("vendor", "area", "vendor area"),
    ("request", "judul_acara", "event title"),
    ("request", "tanggal_acara", "event date"),
    ("request", "lokasi", "event location"),
    ("request", "pengirim_nama", "sender name on the request"),
)


def periksa_konteks(vendor: dict, request: dict) -> list[str]:
    """Labels of the context values that are missing. An empty list means
    buat_spk_docx will not raise over the vendor or the request."""
    sumber = {"vendor": dict(vendor), "request": dict(request)}
    kurang = []
    for asal, kolom, label in KONTEKS_WAJIB:
        nilai = sumber[asal].get(kolom)
        if not (str(nilai).strip() if nilai is not None else ""):
            kurang.append(label)
    return kurang


def _teks(data: dict, kolom: str, sumber: str) -> str:
    """One required text field. Empty is a defect, not a blank to print: an SPK
    with a gap where the location or the scope belongs is worse than no SPK."""
    nilai = data.get(kolom)
    teks = "" if nilai is None else str(nilai).strip()
    if not teks:
        raise ValueError(f"{sumber}.{kolom} is empty — cannot issue the SPK.")
    return teks


def _tanggal(data: dict, kolom: str, sumber: str) -> str:
    """A required date, in Indonesian long form. format_tanggal is the only
    date formatter in the project that produces Indonesian, and the document
    is Indonesian, so it is the one used here."""
    hasil = renderer.format_tanggal(_teks(data, kolom, sumber))
    if not hasil:
        raise ValueError(f"{sumber}.{kolom} is not a usable date.")
    return hasil


def _harga(data: dict) -> int:
    """The amount. Zero is refused along with missing: a work order that
    appoints a vendor for nothing is a filled-in blank, not a value."""
    nilai = data.get("harga")
    if isinstance(nilai, bool) or not isinstance(nilai, int):
        raise ValueError(f"spk.harga must be whole rupiah, got {nilai!r}.")
    if nilai <= 0:
        raise ValueError("spk.harga is zero — cannot issue the SPK.")
    return nilai


def _garis_bawah(paragraf) -> None:
    """Horizontal rule under the letterhead. python-docx has no API for it, so
    the border goes on straight through the underlying XML."""
    pPr = paragraf._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bawah = OxmlElement("w:bottom")
    bawah.set(qn("w:val"), "single")
    bawah.set(qn("w:sz"), "12")
    bawah.set(qn("w:space"), "1")
    bawah.set(qn("w:color"), "auto")
    pBdr.append(bawah)
    pPr.append(pBdr)


def _paragraf(doc, teks="", *, bold=False, align=None, ukuran=None, after=None):
    """One paragraph, with the handful of knobs this document needs."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = after
    if teks:
        run = p.add_run(teks)
        run.bold = bold
        if ukuran is not None:
            run.font.size = ukuran
    return p


def _baris(tabel, label: str, *nilai: str) -> None:
    """One label/value row. Extra values become extra lines in the same cell —
    that is how the terbilang line sits under the amount."""
    sel = tabel.add_row().cells
    sel[0].text = label
    sel[1].text = nilai[0]
    for tambahan in nilai[1:]:
        sel[1].add_paragraph(tambahan)


def buat_spk_docx(spk: dict, vendor: dict, request: dict) -> BytesIO:
    """Render one SPK. Returns a stream positioned at the start; the caller
    decides whether it becomes a download, a file, or an attachment.

    Every field is read before anything is drawn, so a missing value fails
    before a half-built document exists.
    """
    # dict() so a sqlite3.Row passes straight in — Row has no .get().
    spk, vendor, request = dict(spk), dict(vendor), dict(request)

    # Same check the form runs, so the two can never disagree about what a
    # complete SPK needs.
    kurang = periksa_konteks(vendor, request)
    if kurang:
        raise ValueError(f"missing: {', '.join(kurang)} — cannot issue the SPK.")

    nomor = _teks(spk, "nomor", "spk")
    lingkup_kerja = _teks(spk, "lingkup_kerja", "spk")
    termin = _teks(spk, "termin", "spk")
    tanggal_terbit = _tanggal(spk, "tanggal_terbit", "spk")
    harga = _harga(spk)

    nama_pt = _teks(vendor, "nama_pt", "vendor")
    pic_nama = _teks(vendor, "pic_nama", "vendor")
    area = _teks(vendor, "area", "vendor")

    judul_acara = _teks(request, "judul_acara", "request")
    lokasi = _teks(request, "lokasi", "request")
    tanggal_acara = _tanggal(request, "tanggal_acara", "request")
    pengirim_nama = _teks(request, "pengirim_nama", "request")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAMA
    normal.font.size = FONT_UKURAN

    # Letterhead
    _paragraf(doc, KOP_NAMA, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              ukuran=Pt(14), after=Pt(0))
    _paragraf(doc, KOP_ALAMAT, align=WD_ALIGN_PARAGRAPH.CENTER,
              ukuran=Pt(9), after=Pt(0))
    kontak = _paragraf(doc, f"{KOP_TELEPON} · {KOP_EMAIL}",
                       align=WD_ALIGN_PARAGRAPH.CENTER, ukuran=Pt(9), after=Pt(6))
    _garis_bawah(kontak)

    # Title
    _paragraf(doc, "", after=Pt(6))
    _paragraf(doc, "SURAT PERINTAH KERJA", bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, ukuran=Pt(13), after=Pt(0))
    _paragraf(doc, f"No. {nomor}", align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(18))

    # Recipient
    _paragraf(doc, "Kepada Yth.", after=Pt(0))
    _paragraf(doc, nama_pt, bold=True, after=Pt(0))
    _paragraf(doc, f"u.p. {pic_nama}", after=Pt(0))
    _paragraf(doc, area, after=Pt(18))

    _paragraf(
        doc,
        f"Dengan surat ini, {KOP_NAMA} menunjuk {nama_pt} sebagai pelaksana "
        f"pekerjaan untuk acara {judul_acara}, dengan rincian sebagai berikut:",
        after=Pt(12),
    )

    tabel = doc.add_table(rows=0, cols=2)
    tabel.style = "Table Grid"
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _baris(tabel, "Pekerjaan", lingkup_kerja)
    _baris(tabel, "Acara", judul_acara)
    _baris(tabel, "Tanggal acara", tanggal_acara)
    _baris(tabel, "Lokasi", lokasi)
    # The written amount is what a signatory reads against the figure, so it
    # is generated from the same integer rather than typed alongside it.
    _baris(tabel, "Nilai pekerjaan", format_rupiah(harga), f"({terbilang(harga)})")
    _baris(tabel, "Termin", termin)

    for baris in tabel.rows:
        baris.cells[0].width = Cm(4.5)
        baris.cells[1].width = Cm(11.5)

    _paragraf(doc, "", after=Pt(12))
    _paragraf(
        doc,
        "Pekerjaan agar dilaksanakan sesuai rincian di atas dan selesai "
        f"sebelum acara berlangsung pada {tanggal_acara}. Perubahan lingkup "
        "maupun nilai pekerjaan hanya berlaku bila disepakati secara tertulis "
        "oleh kedua belah pihak. Atas kerja samanya kami ucapkan terima kasih.",
        after=Pt(24),
    )

    # Signature block
    kanan = WD_ALIGN_PARAGRAPH.RIGHT
    _paragraf(doc, f"{KOP_KOTA}, {tanggal_terbit}", align=kanan, after=Pt(0))
    _paragraf(doc, KOP_NAMA, align=kanan, after=Pt(0))
    # Room for a wet signature.
    _paragraf(doc, "", after=Pt(0))
    _paragraf(doc, "", after=Pt(0))
    _paragraf(doc, "", after=Pt(0))
    _paragraf(doc, pengirim_nama, bold=True, align=kanan, after=Pt(0))
    _paragraf(doc, "Divisi Procurement", align=kanan)

    aliran = BytesIO()
    doc.save(aliran)
    aliran.seek(0)
    return aliran
